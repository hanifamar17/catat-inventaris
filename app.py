import secrets
from uuid import uuid4
from flask import Flask, Response, jsonify, render_template, request, redirect, send_file, session, url_for, flash, abort
from google.oauth2 import service_account
from googleapiclient.discovery import build
import os
import json
from datetime import datetime, date, timedelta
import calendar
from collections import defaultdict
from flask_wtf.csrf import CSRFProtect
import pytz
import requests
import ssl
from werkzeug.security import check_password_hash
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_session import Session
from flask import send_from_directory
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
import qrcode

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "fallback-if-missing")
csrf = CSRFProtect(app)

# testing (LOCAL ONLY). DO NOT USE IN PRODUCTION!!
#ssl._create_default_https_context = ssl._create_unverified_context


if os.getenv("VERCEL") is None:
    from dotenv import load_dotenv
    load_dotenv()  # Load environment variables from .env file

# Spreadsheet ID dan range untuk menyimpan data
SPREADSHEET_ID = os.getenv("SPREADSHEET_ID")
PDFSHIFT_API_KEY = os.getenv("PDFSHIFT_API_KEY")
SCOPES = ['https://www.googleapis.com/auth/spreadsheets']

# Load credentials
if os.getenv("VERCEL"):
    # Dari environment variable GOOGLE_CREDENTIALS
    service_account_info = json.loads(os.getenv("GOOGLE_CREDENTIALS", "{}"))
    service_account_info['private_key'] = service_account_info['private_key'].replace('\\n', '\n')
    credentials = service_account.Credentials.from_service_account_info(
        service_account_info, scopes=SCOPES
    )
else:
    # Dari file lokal
    credentials = service_account.Credentials.from_service_account_file(
        'inventaris-credentials.json', scopes=SCOPES
    )

sheets_service = build('sheets', 'v4', credentials=credentials)


## OTHERS
#--- Format tanggal ---
@app.template_filter("format_date")
def format_date(value):
    from datetime import datetime
    try:
        dt = datetime.strptime(value, "%Y-%m-%d")
        return dt.strftime("%d/%m/%Y")
    except:
        return value

@app.template_filter('format_date_2')
def format_date_2(value):
    # Pastikan timezone ke Asia/Jakarta
    jakarta = pytz.timezone("Asia/Jakarta")
    if value.tzinfo is None:
        value = jakarta.localize(value)
    else:
        value = value.astimezone(jakarta)

    return value.strftime("%a, %d %b %Y %H.%M")

#--- Format bulan ---
@app.template_filter()
def format_month(value):
    try:
        dt = datetime.strptime(value, "%Y-%m")
        return dt.strftime("%B %Y")  # e.g. July 2025
    except:
        return value
    
@app.template_filter('monthname')
def monthname_filter(m):
    return calendar.month_name[int(m)]

@app.context_processor
def inject_now():
    return {'now': datetime.now}
 

# PDFSHIFT Configuration
def pdf_with_pdfshift(html_content, filename="report.pdf"):
    if not PDFSHIFT_API_KEY:
        print("PDFSHIFT_API_KEY not found in environment variables.")
        return False

    response = requests.post(
        "https://api.pdfshift.io/v3/convert/pdf",
        headers={
            "X-API-Key": PDFSHIFT_API_KEY
        },
        json={
            "source": html_content,
            "landscape": False,
            "use_print": False
        }
    )

    if response.status_code == 200:
        with open(filename, "wb") as f:
            f.write(response.content)
        return True
    else:
        print("PDFShift Error:", response.text)
        return False

#session cookie
app.config['WTF_CSRF_SECRET_KEY'] = os.getenv("WTF_CSRF_SECRET_KEY", "your-random-string")
app.config['SESSION_COOKIE_SECURE'] = True  # Karena di Vercel pakai HTTPS
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

# Login management
login_manager = LoginManager()
login_manager.login_view = 'login'
login_manager.init_app(app)

# username & password admin from googlesheet
def get_accounts_from_sheet():
    sheet = sheets_service.spreadsheets()
    result = sheet.values().get(
        spreadsheetId=SPREADSHEET_ID,
        range="Profil!A2:B"  # Asumsikan header di baris 1
    ).execute()
    
    values = result.get('values', [])
    # Buat dict: {username: password}
    accounts = {row[0]: row[1] for row in values if len(row) >= 2}
    return accounts

class AdminUser(UserMixin):
    def __init__(self, username):
        self.id = username  # Bisa ditampilkan sebagai current_user.id

@login_manager.user_loader
def load_user(user_id):
    return AdminUser(user_id)  # Tidak perlu validasi ulang, cukup restore sesi

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username_input = request.form['username']
        password_input = request.form['password']

        accounts = get_accounts_from_sheet()

        # Autentikasi dari sheet
        if username_input in accounts and check_password_hash(accounts[username_input], password_input):
            user = AdminUser(username_input)
            login_user(user, remember=True)
            return redirect(url_for('dashboard'))
        else:
            flash("Login gagal. Username atau password salah.")
    
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))


#--- Ambil data dari Google Sheets ---
def get_data(sheet_name):
    result = sheets_service.spreadsheets().values().get(
        spreadsheetId=SPREADSHEET_ID,
        range=f"{sheet_name}!A2:G"
    ).execute()

    values = result.get('values', [])
    return values
    
# Home page    
@app.route("/")
def index():
    return redirect(url_for("inventaris"))


# Dashboard page


# generate kode barang
def generate_kode_barang(sheet_name="Barang"):
    data = sheets_service.spreadsheets().values().get(
        spreadsheetId=SPREADSHEET_ID,
        range=f"{sheet_name}!A2:A"
    ).execute().get('values', [])

    if not data:
        return "LAB-001"

    # Ambil kode terakhir yang valid
    kode_terakhir = ""
    for row in reversed(data):
        if row and row[0].startswith("LAB-"):
            kode_terakhir = row[0]
            break

    if kode_terakhir:
        try:
            nomor_terakhir = int(kode_terakhir.replace("LAB-", ""))
        except ValueError:
            nomor_terakhir = 0
    else:
        nomor_terakhir = 0

    kode_baru = f"LAB-{nomor_terakhir + 1:03}"
    return kode_baru

# Income page
@app.route("/inventaris", methods=["GET", "POST"])
#@login_required
def inventaris():
    if request.method == "POST":
        try:
            kode_barang = generate_kode_barang()
            date_inventaris = request.form["date"]
            nama_barang = request.form["nama_barang"]
            merek = request.form["merek"]
            jumlah = request.form["jumlah"]
            kondisi = request.form["kondisi"]
            keterangan = request.form["keterangan"]

            values = [[kode_barang, nama_barang, merek, jumlah, date_inventaris, kondisi, keterangan]]

            sheets_service.spreadsheets().values().append(
                spreadsheetId=SPREADSHEET_ID,
                range="Barang!A2:G",
                valueInputOption="USER_ENTERED",
                body={"values": values}
            ).execute()

            return jsonify({"status": "success"})
        except Exception as e:
            return jsonify({"status": "error", "message": str(e)})
        
    inventaris_data = get_data("Barang")
        
    return render_template("inventaris.html", records=inventaris_data, today=date.today())


# Edit record
@app.route("/edit/<sheet>/<kode_barang>", methods=["POST"])
def edit_record(sheet, kode_barang):
    nama_barang = request.form.get("nama_barang")
    merek = request.form.get("merek")
    jumlah = request.form.get("jumlah")
    date_inventaris = request.form.get("date")
    kondisi = request.form.get("kondisi")
    keterangan = request.form.get("keterangan")

    # Ambil semua data
    data = sheets_service.spreadsheets().values().get(
        spreadsheetId=SPREADSHEET_ID,
        range=f"{sheet.capitalize()}!A2:G"
    ).execute().get('values', [])

    # Cari baris berdasarkan Kode Barang
    for index, row in enumerate(data, start=2):  # mulai dari baris ke-2
        if row[0] == kode_barang:
            values = [[
                kode_barang,
                nama_barang,
                merek,
                jumlah,
                date_inventaris,
                kondisi,
                keterangan
            ]]
            sheets_service.spreadsheets().values().update(
                spreadsheetId=SPREADSHEET_ID,
                range=f"{sheet.capitalize()}!A{index}:G{index}",
                valueInputOption="USER_ENTERED",
                body={"values": values}
            ).execute()
            return jsonify({"status": "success"})

    return jsonify({"status": "error", "message": "Kode Barang tidak ditemukan"})

def get_sheet_data_with_index(sheet_name):
    result = sheets_service.spreadsheets().values().get(
        spreadsheetId=SPREADSHEET_ID,
        range=f"{sheet_name}!A2:G"  # Sesuaikan dengan kolom A-E (ID, Date, Item, Category, Amount)
    ).execute()

    values = result.get('values', [])
    data_with_index = []

    for i, row in enumerate(values, start=2):  # start=2 karena A1 adalah header
        data_with_index.append({
            "index": i,   # Baris aktual di Google Sheets
            "data": row   # Data pada baris tersebut
        })

    return data_with_index

def get_sheet_id_by_name(sheet_name):
    """
    Mengambil sheetId dari nama sheet/tab.
    """
    metadata = sheets_service.spreadsheets().get(spreadsheetId=SPREADSHEET_ID).execute()
    for sheet in metadata["sheets"]:
        if sheet["properties"]["title"].lower() == sheet_name.lower():
            return sheet["properties"]["sheetId"]
    raise ValueError(f"Sheet '{sheet_name}' not found.")

## Delete record
@app.route("/delete/<sheet>/<kode_barang>", methods=["POST"])
def delete_record(sheet, kode_barang):
    try:
        all_data = get_sheet_data_with_index(sheet.capitalize())

        for row in all_data:
            if row["data"][0] == kode_barang:
                row_index = row["index"] - 1  # 0-based index
                break
        else:
            return jsonify({"status": "error", "message": "Kode Barang tidak ditemukan"})

        sheet_id = get_sheet_id_by_name(sheet.capitalize())

        sheets_service.spreadsheets().batchUpdate(
            spreadsheetId=SPREADSHEET_ID,
            body={
                "requests": [
                    {
                        "deleteDimension": {
                            "range": {
                                "sheetId": sheet_id,
                                "dimension": "ROWS",
                                "startIndex": row_index,
                                "endIndex": row_index + 1
                            }
                        }
                    }
                ]
            }
        ).execute()

        return jsonify({"status": "success", "message": "Barang berhasil dihapus"})
    except Exception as e:
        return jsonify({"status": "error", "message": f"Gagal menghapus barang: {e}"})

## Cetak label barang
def get_barang_by_kode(kode_barang):
    # Ambil semua data dari sheet
    result = sheets_service.spreadsheets().values().get(
        spreadsheetId=SPREADSHEET_ID,
        range='Barang!A2:G'
    ).execute()

    values = result.get('values', [])

    for row in values:
        if row[0] == kode_barang:
            return {
                'kode_barang': row[0],
                'nama_barang': row[1],
                'merek': row[2],
                'kondisi': row[5]
            }

    return None

@app.route('/cetak-label/<kode_barang>')
def cetak_label(kode_barang):
    # Misal ambil data dari Google Sheets atau DB
    barang = get_barang_by_kode(kode_barang)  # buat fungsi ini sesuai datamu

    if not barang:
        return "Barang tidak ditemukan", 404

    nama_barang = barang['nama_barang']  # pastikan key sesuai
    merek = barang['merek']
    kondisi = barang['kondisi']

    # Generate QR Code
    qr = qrcode.make(kode_barang)
    qr_io = BytesIO()
    qr.save(qr_io, format='PNG')
    qr_io.seek(0)

    # Buat dokumen Word
    doc = Document()

    # Buat table 1 kolom, 3 baris (QR + Text)
    table = doc.add_table(rows=3, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    table.style = 'Table Grid'

    # Tambahkan border
    #tbl = table._tbl
    #tbl.set(qn('w:tblBorders'), 
    #    '''
    #    <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    #        <w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>
    #        <w:left w:val="single" w:sz="6" w:space="0" w:color="000000"/>
    #        <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>
    #        <w:right w:val="single" w:sz="6" w:space="0" w:color="000000"/>
    #    </w:tblBorders>
    #    '''
    #)

    # Baris 1: Judul label
    cell0 = table.cell(0, 0)
    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run0 = p0.add_run("Label Inventaris")
    run0.bold = True
    run0.font.size = Pt(12)

    # Baris 2: QR code
    cell1 = table.cell(1, 0)
    p1 = cell1.paragraphs[0]
    run1 = p1.add_run()
    run1.add_picture(qr_io, width=Inches(1.5))
    p1.alignment = 1  # Center

    # Baris 3: Text
    cell2 = table.cell(2, 0)
    p2 = cell2.paragraphs[0]
    p2.alignment = 1  # Center
    run2 = p2.add_run(f"{kode_barang}\n{nama_barang}\n{merek}\n{kondisi}")
    run2.font.size = Pt(10)

    # Simpan ke memory
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    return send_file(
        doc_io,
        as_attachment=True,
        download_name=f'label_{kode_barang}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

## Unduh annual report pdf
#@app.route('/annual_report')
#@login_required
#def annual_report():
#    selected_year = request.args.get("year", datetime.now().year, type=int)
#    # Ambil data dari Google Sheets
#    income_data = get_data("Income")
#    expenses_data = get_data("Expenses")
#
#    # Buat struktur penampung total per bulan
#    monthly_summary = defaultdict(lambda: {"income": 0, "expenses": 0})
#
#    def process_data(data, tipe):
#        for row in data:
#            try:
#                date_obj = datetime.strptime(row[1], "%Y-%m-%d")
#                if date_obj.year == selected_year:
#                    month_key = date_obj.strftime("%Y-%m")
#                    amount = int(row[4].replace(".", "").replace(",", ""))
#                    monthly_summary[month_key][tipe] += amount
#            except (IndexError, ValueError):
#                continue  # Lewati baris rusak
#
#    process_data(income_data, "income")
#    process_data(expenses_data, "expenses")
#
#    # Tambahkan saldo
#    for key in monthly_summary:
#        summary = monthly_summary[key]
#        summary["balance"] = summary["income"] - summary["expenses"]
#
#    # Urutkan berdasarkan bulan
#    sorted_months = sorted(monthly_summary.keys())
#    summary_data = [(month, monthly_summary[month]) for month in sorted_months]
#
#    chart_labels = [calendar.month_name[m] for m in range(1, 13)]
#    chart_income = []
#    chart_expenses = []
#    chart_balance = []
#
#    for m in range(1, 13):
#        key = f"{selected_year}-{m:02d}"
#        data = monthly_summary.get(key, {"income": 0, "expenses": 0, "balance": 0})
#        chart_income.append(data["income"])
#        chart_expenses.append(data["expenses"])
#        chart_balance.append(data["balance"])
#    
#    total_income = sum(chart_income)
#    total_expenses = sum(chart_expenses)
#    total_balance = sum(chart_balance)
#
#    # growth bulan ini
#    net_balance = [i - e for i, e in zip(chart_income, chart_expenses)]
#    income_growth = get_growth(chart_income)
#    expenses_growth = get_growth(chart_expenses)
#    net_growth = get_growth(net_balance)
#
#    # growth series satu tahun
#    income_growth_series = get_growth_series(chart_income)
#    expenses_growth_series = get_growth_series(chart_expenses)
#    net_growth_series = get_growth_series(net_balance)
#
#    now = datetime.now()
#
#    html = render_template("report.html",
#                           now=now,
#                           summary_data=summary_data,
#                           selected_year=str(selected_year),
#                           chart_labels=chart_labels,
#                           chart_income=chart_income,
#                           chart_expenses=chart_expenses,
#                           chart_balance=chart_balance,
#                           total_income=total_income,
#                           total_expenses=total_expenses,
#                           total_balance=total_balance,
#                           income_growth=income_growth,
#                           expenses_growth=expenses_growth,
#                           net_growth=net_growth,
#                           income_growth_series=income_growth_series,
#                           expenses_growth_series=expenses_growth_series,
#                           net_growth_series=net_growth_series)  # render HTML dengan Jinja
#    
#    response = requests.post(
#        "https://api.pdfshift.io/v3/convert/pdf",
#        headers={ "X-API-Key": PDFSHIFT_API_KEY },
#        json={ "source": html }
#    )
#
#    if response.status_code == 200:
#        return Response(
#            response.content,
#            mimetype='application/pdf',
#            headers={
#                "Content-Disposition": "inline; filename=report.pdf"
#            }
#        )
#    else:
#        return f"PDFShift Error: {response.text}", 500
#
## Unduh report pdf (use for testing only)
#@app.route('/annual-report-test')
#@login_required
#def annual_report_test():
#    selected_year = request.args.get("year", datetime.now().year, type=int)
#    
#    # Ambil data dari Google Sheets
#    income_data = get_data("Income")
#    expenses_data = get_data("Expenses")
#
#    # Buat struktur penampung total per bulan
#    monthly_summary = defaultdict(lambda: {"income": 0, "expenses": 0})
#
#    def process_data(data, tipe):
#        for row in data:
#            try:
#                date_obj = datetime.strptime(row[1], "%Y-%m-%d")
#                if date_obj.year == selected_year:
#                    month_key = date_obj.strftime("%Y-%m")
#                    amount = int(row[4].replace(".", "").replace(",", ""))
#                    monthly_summary[month_key][tipe] += amount
#            except (IndexError, ValueError):
#                continue  # Lewati baris rusak
#
#    process_data(income_data, "income")
#    process_data(expenses_data, "expenses")
#
#    # Tambahkan saldo
#    for key in monthly_summary:
#        summary = monthly_summary[key]
#        summary["balance"] = summary["income"] - summary["expenses"]
#
#    # Urutkan berdasarkan bulan
#    sorted_months = sorted(monthly_summary.keys())
#    summary_data = [(month, monthly_summary[month]) for month in sorted_months]
#
#    chart_labels = [calendar.month_name[m] for m in range(1, 13)]
#    chart_income = []
#    chart_expenses = []
#    chart_balance = []
#
#    for m in range(1, 13):
#        key = f"{selected_year}-{m:02d}"
#        data = monthly_summary.get(key, {"income": 0, "expenses": 0, "balance": 0})
#        chart_income.append(data["income"])
#        chart_expenses.append(data["expenses"])
#        chart_balance.append(data["balance"])
#
#    total_income = sum(chart_income)
#    total_expenses = sum(chart_expenses)
#    total_balance = sum(chart_balance)
#
#    # growth bulan ini
#    net_balance = [i - e for i, e in zip(chart_income, chart_expenses)]
#    income_growth = get_growth(chart_income)
#    expenses_growth = get_growth(chart_expenses)
#    net_growth = get_growth(net_balance)
#
#    # growth series satu tahun
#    income_growth_series = get_growth_series(chart_income)
#    expenses_growth_series = get_growth_series(chart_expenses)
#    net_growth_series = get_growth_series(net_balance)
#
#    now = datetime.now()
#
#    return render_template("report.html", 
#                           now = now,
#                           summary_data=summary_data,
#                           selected_year=str(selected_year),
#                           chart_labels=chart_labels,
#                           chart_income=chart_income,
#                           chart_expenses=chart_expenses,
#                           chart_balance=chart_balance,
#                           total_income=total_income,
#                           total_expenses=total_expenses,
#                           total_balance=total_balance,
#                           income_growth=income_growth,
#                           expenses_growth=expenses_growth,
#                           net_growth=net_growth,
#                           income_growth_series=income_growth_series,
#                           expenses_growth_series=expenses_growth_series,
#                           net_growth_series=net_growth_series)
#
#@app.route("/monthly_report/<month>")
#@login_required
#def monthly_report(month):
#    # Ambil data dari Google Sheets
#    income_data = get_data("Income")
#    expense_data = get_data("Expenses")
#
#    try:
#        year, mon = map(int, month.split('-'))
#    except ValueError:
#        return "Invalid month format. Use YYYY-MM.", 400
#
#    def filter_by_month(data, y, m):
#        filtered = []
#        for row in data:
#            try:
#                row_date = datetime.strptime(row[1], "%Y-%m-%d")
#                if row_date.month == m and row_date.year == y:
#                    filtered.append(row)
#            except:
#                continue
#        return filtered
#
#    def get_total(data):
#        total = 0
#        for row in data:
#            try:
#                amount = float(row[4].replace(",", "").replace(".", ""))
#                total += amount
#            except:
#                continue
#        return total
#
#    # Data bulan ini
#    income_filtered = filter_by_month(income_data, year, mon)
#    expense_filtered = filter_by_month(expense_data, year, mon)
#
#    income_totals = defaultdict(float)
#    for row in income_filtered:
#        try:
#            category = row[3]
#            amount = float(row[4].replace(",", "").replace(".", ""))
#            income_totals[category] += amount
#        except:
#            continue
#
#    expense_totals = defaultdict(float)
#    for row in expense_filtered:
#        try:
#            category = row[3]
#            amount = float(row[4].replace(",", "").replace(".", ""))
#            expense_totals[category] += amount
#        except:
#            continue
#
#    total_income = sum(income_totals.values())
#    total_expense = sum(expense_totals.values())
#    balance = total_income - total_expense
#
#    # Data bulan sebelumnya
#    prev_year = year if mon > 1 else year - 1
#    prev_month = mon - 1 if mon > 1 else 12
#
#    prev_income_filtered = filter_by_month(income_data, prev_year, prev_month)
#    prev_expense_filtered = filter_by_month(expense_data, prev_year, prev_month)
#
#    prev_total_income = get_total(prev_income_filtered)
#    prev_total_expense = get_total(prev_expense_filtered)
#    prev_balance = prev_total_income - prev_total_expense
#
#    # Growth perbandingan bulan ini vs sebelumnya
#    income_growth = get_growth([prev_total_income, total_income])
#    expenses_growth = get_growth([prev_total_expense, total_expense])
#    net_growth = get_growth([prev_balance, balance])
#
#    # Nama bulan
#    month_name = calendar.month_name[mon]
#    now = datetime.now()
#
#    html = render_template("monthly-report.html",
#                           now=now,
#                           month=month,
#                           selected_year=year,
#                           month_name=month_name,
#                           income_labels=list(income_totals.keys()),
#                           income_data=list(income_totals.values()),
#                           expense_labels=list(expense_totals.keys()),
#                           expense_data=list(expense_totals.values()),
#                           total_income=total_income,
#                           total_expenses=total_expense,
#                           balance=balance,
#                           income_growth=income_growth,
#                           expenses_growth=expenses_growth,
#                           net_growth=net_growth,
#                           summary_data_income=income_filtered,
#                           summary_data_expenses=expense_filtered
#    )
#
#    response = requests.post(
#        "https://api.pdfshift.io/v3/convert/pdf",
#        headers={ "X-API-Key": PDFSHIFT_API_KEY },
#        json={ "source": html }
#    )
#
#    if response.status_code == 200:
#        return Response(
#            response.content,
#            mimetype='application/pdf',
#            headers={
#                "Content-Disposition": "inline; filename=report.pdf"
#            }
#        )
#    else:
#        return f"PDFShift Error: {response.text}", 500
#
## Unduh report perbulan pdf (use for testing only)
#@app.route("/monthly-report-tes/<month>")
#@login_required
#def monthly_report_test(month):
#    # Ambil data dari Google Sheets
#    income_data = get_data("Income")
#    expense_data = get_data("Expenses")
#
#    try:
#        year, mon = map(int, month.split('-'))
#    except ValueError:
#        return "Invalid month format. Use YYYY-MM.", 400
#
#    def filter_by_month(data, y, m):
#        filtered = []
#        for row in data:
#            try:
#                row_date = datetime.strptime(row[1], "%Y-%m-%d")
#                if row_date.month == m and row_date.year == y:
#                    filtered.append(row)
#            except:
#                continue
#        return filtered
#
#    def get_total(data):
#        total = 0
#        for row in data:
#            try:
#                amount = float(row[4].replace(",", "").replace(".", ""))
#                total += amount
#            except:
#                continue
#        return total
#
#    # Data bulan ini
#    income_filtered = filter_by_month(income_data, year, mon)
#    expense_filtered = filter_by_month(expense_data, year, mon)
#
#    income_totals = defaultdict(float)
#    for row in income_filtered:
#        try:
#            category = row[3]
#            amount = float(row[4].replace(",", "").replace(".", ""))
#            income_totals[category] += amount
#        except:
#            continue
#
#    expense_totals = defaultdict(float)
#    for row in expense_filtered:
#        try:
#            category = row[3]
#            amount = float(row[4].replace(",", "").replace(".", ""))
#            expense_totals[category] += amount
#        except:
#            continue
#
#    total_income = sum(income_totals.values())
#    total_expense = sum(expense_totals.values())
#    balance = total_income - total_expense
#
#    # Data bulan sebelumnya
#    prev_year = year if mon > 1 else year - 1
#    prev_month = mon - 1 if mon > 1 else 12
#
#    prev_income_filtered = filter_by_month(income_data, prev_year, prev_month)
#    prev_expense_filtered = filter_by_month(expense_data, prev_year, prev_month)
#
#    prev_total_income = get_total(prev_income_filtered)
#    prev_total_expense = get_total(prev_expense_filtered)
#    prev_balance = prev_total_income - prev_total_expense
#
#    # Growth perbandingan bulan ini vs sebelumnya
#    income_growth = get_growth([prev_total_income, total_income])
#    expenses_growth = get_growth([prev_total_expense, total_expense])
#    net_growth = get_growth([prev_balance, balance])
#
#    # Nama bulan
#    month_name = calendar.month_name[mon]
#    now = datetime.now()
#
#    return render_template("monthly-report.html",
#                           now=now,
#                           month=month,
#                           selected_year=year,
#                           month_name=month_name,
#                           income_labels=list(income_totals.keys()),
#                           income_data=list(income_totals.values()),
#                           expense_labels=list(expense_totals.keys()),
#                           expense_data=list(expense_totals.values()),
#                           total_income=total_income,
#                           total_expenses=total_expense,
#                           balance=balance,
#                           income_growth=income_growth,
#                           expenses_growth=expenses_growth,
#                           net_growth=net_growth,
#                           summary_data_income=income_filtered,
#                           summary_data_expenses=expense_filtered
#    )




if __name__ == '__main__':
    app.run(debug=True)
